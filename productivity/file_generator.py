import os
import uuid
from docx import Document
from openpyxl import Workbook

def create_text_file(filename, content, output_folder="static/generated"):
    os.makedirs(output_folder, exist_ok=True)
    safe_name = filename.replace(" ", "_")
    path = os.path.join(output_folder, safe_name)
    with open(path, "w", encoding="utf-8") as file:
        file.write(content)
    return path

def create_docx(title, content, output_folder="static/generated"):
    os.makedirs(output_folder, exist_ok=True)
    filename = f"saivex_document_{uuid.uuid4().hex[:8]}.docx"
    path = os.path.join(output_folder, filename)
    document = Document()
    document.add_heading(title, level=1)
    for line in content.split("\n"):
        document.add_paragraph(line)
    document.save(path)
    return path

def create_excel(rows, output_folder="static/generated"):
    os.makedirs(output_folder, exist_ok=True)
    filename = f"saivex_sheet_{uuid.uuid4().hex[:8]}.xlsx"
    path = os.path.join(output_folder, filename)
    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(row)
    workbook.save(path)
    return path
